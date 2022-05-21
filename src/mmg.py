"""
created by Brad Hillier - BradleyHillier@icloud.com
https://github.com/BradHillier/morning-meeting-generator

generate the morning meeting self.document for Sealegs Kayaking Adventures 
in Ladysmith, BC
"""

from datetime import datetime
from docx import Document
import os.path
import sys

from Tide import get_tides
from Weather import get_weather
from Booking import get_bookings
from SafetyTopics import SafetyTopics 

import config


class MeetingDocumentGenerator:
    def __init__(self):
        config.load()
        self.document = Document()
        self.tides: dict[str, list[Tide]] = get_tides()
        self.weather: list[Weather] = get_weather(10,18)
        self.bookings: list[Booking] = get_bookings(config.CONFIG['calendar ID'])
        self.safety_topics = SafetyTopics(self)

    def generate(self):
        datestring = datetime.now().strftime('%A - %B %d - %Y')
        self.document.add_heading('Daily Safety Meeting', 0)
        self.document.add_paragraph(datestring)

        self._add_attendants()
        self._add_safety_topic(self.weather, self.tides)
        
        self.document.add_heading('General Notes/Maintenance', 2)
        self.document.add_paragraph('\n' * 2)

        self._add_tides()
        self._add_weather()
        self._add_bookings()

    def write_to_file(self):
        try:
            self.document.save(config.CONFIG['output location'] + 'morning_meeting_' + 
                          datetime.now().strftime('%d-%m-%y') + '.docx')
        except FileNotFoundError:
            print('invalid output location, path to directory does not exist')
            sys.exit()
        print('self.document successfully created')

    def _add_attendants(self):
        self.document.add_heading('Attendants', 2)
        p = self.document.add_paragraph()
        for name in config.CONFIG['employees']:
            p.add_run(f'\u2751 {name}         ')

    def _add_safety_topic(self, weather=None, tides=None):
         self.document.add_heading('Safety Topic', 2)
         self.document.add_paragraph('\n'.join(self.safety_topics.topics))

    def _add_tides(self):
        self.document.add_heading('Tides', 2)
        tides_table = self.document.add_table(rows=2, cols=4)
        for idx, tide in enumerate(self.tides['high and low']):
            column_cells = tides_table.column_cells(idx)
            column_cells[0].text = tide.time.strftime('%-I:%M %p')
            column_cells[1].text = f'{tide.meters} meters'


    def _add_weather(self):
        self.document.add_heading('Weather', 2)
        wx_table = self.document.add_table(rows=4, cols=len(self.weather))
        for i in range(len(wx_table.columns)):
            col = wx_table.columns[i]
            col.cells[0].text = self.weather[i].date.strftime('%-I%p').lower()
            col.cells[1].text = self.weather[i].emoji
            col.cells[2].text = f'{self.weather[i].temp}\N{DEGREE SIGN} C',
            col.cells[3].text = self.weather[i].wind

    def _add_bookings(self):
        self.document.add_heading('Bookings', 2)
        bookings_table = self.document.add_table(rows=0, cols=2)
        for booking in self.bookings:
            row_cells = bookings_table.add_row().cells
            start = booking.start_at.strftime('%-I:%M %p')
            end = booking.end_at.strftime('%-I:%M %p')
            row_cells[0].text = f'{start} to {end}'
            row_cells[1].text = booking.title
